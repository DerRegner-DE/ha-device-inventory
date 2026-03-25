# -*- coding: utf-8 -*-
"""
Erstellt alle 3 Dokumentationen für Geräteverwaltung v1.3.5
- Benutzerhandbuch
- Installationsanleitung
- Testplan
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

# ============================================================
# Helper Functions
# ============================================================

def set_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    for level, size in [(1, 18), (2, 14), (3, 12)]:
        h = doc.styles[f'Heading {level}']
        h.font.name = 'Calibri'
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = RGBColor(0x1a, 0x56, 0x8e)

def add_title_page(doc, title, subtitle, date="25.03.2026", author="Matthias Regner"):
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x56, 0x8e)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 50)
    run.font.color.rgb = RGBColor(0x1a, 0x56, 0x8e)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Version 1.3.5  |  Stand: {date}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(author)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    doc.add_page_break()

def _shade_cell(cell, color):
    shading = cell._element.get_or_add_tcPr()
    elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color
    })
    shading.append(elem)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(cell, '1A568E')
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
            if r_idx % 2 == 0:
                _shade_cell(cell, 'F0F5FA')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_tip(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("Tipp: ")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0d, 0x7a, 0x3e)
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0x0d, 0x7a, 0x3e)
    run.font.italic = True

def add_warning(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("Hinweis: ")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
    run.font.italic = True

def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix + " ")
        run.font.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

# ============================================================
# 1. BENUTZERHANDBUCH
# ============================================================

def create_benutzerhandbuch():
    doc = Document()
    set_style(doc)
    add_title_page(doc, "Geräteverwaltung", "Benutzerhandbuch\nAnleitung zur Inventarisierung von Smart-Home-Geräten")

    # --- 1. Einführung ---
    doc.add_heading('1. Was ist die Geräteverwaltung?', level=1)
    doc.add_paragraph(
        'Die Geräteverwaltung ist ein Home Assistant Add-on zur systematischen '
        'Erfassung und Verwaltung aller Geräte im Smart Home. Jedes Gerät '
        '(Router, Steckdose, Kamera, Sensor, etc.) bekommt einen Eintrag mit '
        'allen wichtigen Informationen.'
    )

    doc.add_heading('Funktionsübersicht', level=2)
    add_table(doc, ['Funktion', 'Free', 'Pro'], [
        ['Geräte erfassen (Name, Typ, Standort, Modell)', 'Ja (max. 50)', 'Unbegrenzt'],
        ['Geräte suchen, filtern, Dashboard', 'Ja', 'Ja'],
        ['Offline-Modus (Daten lokal gespeichert)', 'Ja', 'Ja'],
        ['Fotos aufnehmen (Typenschild, Standort)', '–', 'Ja'],
        ['QR-Code / Barcode scannen mit Feld-Befüllung', '–', 'Ja'],
        ['Excel-Export', '–', 'Ja'],
        ['5 Sprachen (DE, EN, ES, FR, RU)', '–', 'Ja'],
    ], col_widths=[8, 3, 3])

    add_tip(doc, 'Die Pro-Lizenz kostet einmalig 9,99 € und schaltet alle Funktionen dauerhaft frei.')

    # --- 2. App öffnen ---
    doc.add_heading('2. App öffnen', level=1)
    doc.add_paragraph(
        'Die Geräteverwaltung ist als Home Assistant Add-on installiert und '
        'direkt über die HA-Seitenleiste erreichbar.'
    )

    doc.add_heading('Über die HA-Seitenleiste (empfohlen)', level=2)
    add_bullet(doc, 'Home Assistant öffnen (Browser oder App)')
    add_bullet(doc, 'In der Seitenleiste auf „Geräteverwaltung" klicken')
    add_bullet(doc, 'Die App öffnet sich direkt im HA-Interface (Ingress)')

    doc.add_heading('Über das Handy', level=2)
    doc.add_paragraph(
        'Die App kann auch direkt auf dem Handy geöffnet werden, '
        'entweder über die HA Companion App (Seitenleiste) oder '
        'im Browser über die HA-URL.'
    )

    doc.add_heading('Verbindungs-Anzeige', level=2)
    add_table(doc, ['Anzeige', 'Bedeutung'], [
        ['Grüner Punkt', 'Verbindung zum Server steht – alles ok!'],
        ['Gelber Punkt', 'Daten warten auf Synchronisation'],
        ['Roter Punkt', 'Offline-Modus – Daten werden lokal gespeichert'],
    ], col_widths=[4, 10])

    add_tip(doc, 'Im Offline-Modus (roter Punkt) können Geräte trotzdem angelegt und bearbeitet werden. Die Daten werden automatisch synchronisiert, sobald wieder eine Verbindung besteht.')

    # --- 3. Sprache wechseln ---
    doc.add_heading('3. Sprache wechseln', level=1)
    doc.add_paragraph(
        'Die App unterstützt 5 Sprachen. Die Spracheinstellung wird in den '
        'Add-on-Einstellungen festgelegt.'
    )
    add_bullet(doc, 'Home Assistant → Einstellungen → Add-ons → Geräteverwaltung')
    add_bullet(doc, 'Tab „Konfiguration" öffnen')
    add_bullet(doc, 'Sprache wählen: de (Deutsch), en (Englisch), es (Spanisch), fr (Französisch), ru (Russisch)')
    add_bullet(doc, 'Speichern und Add-on neustarten')

    add_warning(doc, 'In der Free-Version ist nur Englisch verfügbar. Alle 5 Sprachen erfordern eine Pro-Lizenz.')

    # --- 4. Dashboard ---
    doc.add_heading('4. Dashboard', level=1)
    doc.add_paragraph(
        'Das Dashboard ist die Startseite der App. Hier sieht man auf einen Blick:'
    )
    add_bullet(doc, 'Wie viele Geräte insgesamt erfasst sind')
    add_bullet(doc, 'Auflistung nach Gerätetyp (Router, Steckdosen, Kameras, …)')
    add_bullet(doc, 'Auflistung nach Standort (welcher Raum, welche Etage)')
    add_bullet(doc, 'Schnellzugriff auf häufig verwendete Funktionen')

    # --- 5. Gerät hinzufügen ---
    doc.add_heading('5. Gerät hinzufügen', level=1)
    doc.add_paragraph('So wird ein neues Gerät in der App erfasst:')

    doc.add_heading('Schritt 1: Auf „+" tippen', level=2)
    doc.add_paragraph('In der unteren Leiste auf das Plus-Zeichen tippen.')

    doc.add_heading('Schritt 2: Gerätetyp wählen', level=2)
    doc.add_paragraph(
        'Den passenden Typ auswählen: Router, Switch, Access Point, '
        'Steckdose, Kamera, Sensor, Thermostat, Licht, Sonstiges.'
    )

    doc.add_heading('Schritt 3: Bezeichnung eingeben', level=2)
    doc.add_paragraph(
        'Einen eindeutigen Namen vergeben, z.\u202fB. „FritzBox 6660 Cable" oder „Steckdose Küche".'
    )

    doc.add_heading('Schritt 4: Standort wählen', level=2)
    doc.add_paragraph(
        'Den Raum auswählen, in dem das Gerät steht. Die Räume sind nach Etage gruppiert '
        '(OG, UG, Garten, 19a OG, 19a UG, Garagen, etc.).'
    )

    doc.add_heading('Schritt 5: Foto aufnehmen (Pro)', level=2)
    doc.add_paragraph('Die Kamera-Funktion bietet zwei Optionen:')
    add_bullet(doc, '„Take Photo" – Öffnet die Dateiauswahl bzw. die Handy-Kamera direkt', bold_prefix='Take Photo:')
    add_bullet(doc, '„Use Webcam" – Startet den Live-Kamera-Stream mit Vorschau und Auslöser', bold_prefix='Use Webcam:')
    doc.add_paragraph(
        'Die aufgenommenen Fotos werden als Vorschau angezeigt. '
        'Mit „Bestätigen" wird das Foto dem Gerät zugeordnet, '
        'mit „Erneut aufnehmen" kann ein neues Foto gemacht werden.'
    )
    add_tip(doc, 'Auf dem Handy funktioniert „Take Photo" am zuverlässigsten, '
        'da es die native Kamera-App nutzt. „Use Webcam" benötigt getUserMedia-Zugriff, '
        'der in manchen Browsern eingeschränkt sein kann.')

    doc.add_heading('Schritt 6: QR-Code / Barcode scannen (Pro)', level=2)
    doc.add_paragraph(
        'Der QR-Scanner erkennt Codes und füllt automatisch die passenden Felder aus:'
    )
    add_bullet(doc, 'Seriennummern werden ins Seriennummer-Feld eingetragen')
    add_bullet(doc, 'MAC-Adressen (Format AA:BB:CC:DD:EE:FF) werden erkannt')
    add_bullet(doc, 'URLs und Texte werden im Anmerkungen-Feld gespeichert')
    add_bullet(doc, 'Die aufklappbaren Formular-Abschnitte öffnen sich automatisch, wenn ein Feld befüllt wird')

    add_tip(doc, 'Für beste Ergebnisse den Barcode bei guter Beleuchtung scannen und das Handy ruhig halten.')

    doc.add_heading('Schritt 7: Weitere Details ausfüllen', level=2)
    doc.add_paragraph('Die aufklappbaren Abschnitte enthalten zusätzliche Felder:')
    add_bullet(doc, 'Hersteller, Modell, Seriennummer')
    add_bullet(doc, 'IP-Adresse, MAC-Adresse')
    add_bullet(doc, 'Kaufdatum, Notizen')
    add_bullet(doc, 'Netzwerktyp (WLAN, LAN, Zigbee, Bluetooth, DECT)')
    add_bullet(doc, 'Stromversorgung (Netzteil, PoE, Batterie, USB)')

    doc.add_heading('Schritt 8: Speichern', level=2)
    doc.add_paragraph(
        '„Speichern" am Ende des Formulars antippen. Das Gerät erscheint jetzt in der Geräteliste.'
    )

    add_tip(doc, 'Nicht alle Felder müssen ausgefüllt werden! Name, Typ und Standort reichen für den Anfang. Details können später ergänzt werden.')

    # --- 6. Geräte suchen ---
    doc.add_heading('6. Geräte suchen und filtern', level=1)
    add_bullet(doc, 'In der unteren Leiste auf „Geräte" tippen')
    add_bullet(doc, 'Im Suchfeld oben einen Suchbegriff eingeben')
    add_bullet(doc, 'Die Suche durchsucht alle Felder: Name, Modell, Standort, Seriennummer, etc.')
    doc.add_paragraph(
        'Unter dem Suchfeld befinden sich Filter-Chips. Damit kann die Liste nach '
        'Gerätetyp gefiltert werden (Alle, Router, Steckdosen, Kameras, Sensoren, …).'
    )

    # --- 7. Bearbeiten/Löschen ---
    doc.add_heading('7. Gerät bearbeiten oder löschen', level=1)
    add_bullet(doc, 'In der Geräteliste auf das gewünschte Gerät tippen')
    add_bullet(doc, 'Die Detailansicht öffnet sich mit allen Informationen')
    add_bullet(doc, '„Bearbeiten" antippen, Änderungen vornehmen, „Speichern" tippen')
    add_bullet(doc, 'Zum Löschen: „Löschen" antippen und bestätigen')

    # --- 8. Offline-Modus ---
    doc.add_heading('8. Offline-Modus', level=1)
    doc.add_paragraph(
        'Die App funktioniert auch ohne Netzwerkverbindung. Das ist besonders '
        'nützlich im Keller, in der Garage oder im Garten.'
    )
    add_table(doc, ['Status', 'Was passiert?'], [
        ['Grün', 'Online – Daten werden sofort gespeichert'],
        ['Gelb', 'Änderungen warten auf Synchronisation'],
        ['Rot', 'Offline – Daten werden lokal gespeichert (IndexedDB)'],
    ], col_widths=[3, 11])

    add_warning(doc, 'Fotos können im Offline-Modus aufgenommen werden, werden aber erst bei bestehender Verbindung auf den Server hochgeladen.')

    # --- 9. Excel-Export ---
    doc.add_heading('9. Excel-Export (Pro)', level=1)
    add_bullet(doc, 'Auf „Einstellungen" in der unteren Leiste tippen')
    add_bullet(doc, '„Excel exportieren" antippen')
    add_bullet(doc, 'Die Datei wird erstellt und zum Download angeboten')
    doc.add_paragraph(
        'Die Excel-Datei enthält alle Geräte mit sämtlichen erfassten Feldern.'
    )

    # --- 10. Lizenz ---
    doc.add_heading('10. Lizenz-System', level=1)
    doc.add_paragraph('Die Geräteverwaltung gibt es in zwei Versionen:')

    doc.add_heading('Free-Version', level=2)
    add_bullet(doc, 'Maximal 50 Geräte')
    add_bullet(doc, 'Nur Englisch')
    add_bullet(doc, 'Kein Kamera/Scanner/Export')

    doc.add_heading('Pro-Version (9,99 € einmalig)', level=2)
    add_bullet(doc, 'Unbegrenzte Geräte')
    add_bullet(doc, '5 Sprachen (DE, EN, ES, FR, RU)')
    add_bullet(doc, 'Kamera (Foto-Aufnahme)')
    add_bullet(doc, 'Barcode/QR-Scanner mit Feld-Befüllung')
    add_bullet(doc, 'Excel-Export')

    doc.add_heading('Lizenz aktivieren', level=2)
    add_bullet(doc, 'Home Assistant → Einstellungen → Add-ons → Geräteverwaltung')
    add_bullet(doc, 'Tab „Konfiguration"')
    add_bullet(doc, 'Lizenzschlüssel im Feld „license_key" eintragen')
    add_bullet(doc, 'Speichern – die Lizenz wird sofort aktiviert')

    add_tip(doc, 'Die Lizenz wird serverseitig gespeichert und bleibt auch nach Add-on-Updates erhalten.')

    # --- 11. Tipps ---
    doc.add_heading('11. Tipps für die Inventarisierung', level=1)
    add_bullet(doc, 'Raum für Raum vorgehen – systematisch einen Raum nach dem anderen durchgehen', bold_prefix='Systematisch:')
    add_bullet(doc, 'Typenschild immer fotografieren – Modell, SN und MAC-Adresse sind dort zu finden', bold_prefix='Typenschild:')
    add_bullet(doc, 'MAC-Adresse auf der Rückseite/Unterseite suchen (Format: AA:BB:CC:DD:EE:FF)', bold_prefix='MAC-Adresse:')
    add_bullet(doc, 'IP-Adresse notieren – steht oft im Geräte-Display oder in der FritzBox', bold_prefix='IP-Adresse:')
    add_bullet(doc, 'Stromversorgung und Netzwerkart erfassen – hilft bei der Fehlersuche', bold_prefix='Details:')

    doc.save('Benutzerhandbuch_Geraeteverwaltung.docx')
    print("Benutzerhandbuch erstellt.")

# ============================================================
# 2. INSTALLATIONSANLEITUNG
# ============================================================

def create_installationsanleitung():
    doc = Document()
    set_style(doc)
    add_title_page(doc, "Geräteverwaltung", "Installationsanleitung\nInstallation als Home Assistant Add-on")

    # --- 1. Übersicht ---
    doc.add_heading('1. Übersicht', level=1)
    doc.add_paragraph(
        'Die Geräteverwaltung wird als Home Assistant Add-on installiert. '
        'Das Add-on beinhaltet Frontend (Preact + Tailwind), Backend (FastAPI + SQLite) '
        'und wird als Docker-Container betrieben. Die Installation erfolgt '
        'vollständig über die HA-Oberfläche.'
    )

    doc.add_heading('Architektur', level=2)
    add_table(doc, ['Komponente', 'Technologie', 'Beschreibung'], [
        ['Frontend', 'Preact + TypeScript + Tailwind', 'PWA mit Offline-Fähigkeit (IndexedDB)'],
        ['Backend', 'Python FastAPI + SQLite', 'REST API auf Port 3001 (intern)'],
        ['Container', 'Docker (Alpine Linux)', 'Nginx Reverse Proxy + Uvicorn'],
        ['Zugriff', 'HA Ingress', 'Direkt in der HA-Seitenleiste'],
    ], col_widths=[3, 5, 6])

    doc.add_heading('Voraussetzungen', level=2)
    add_bullet(doc, 'Home Assistant OS oder Supervised (mit Add-on-Support)')
    add_bullet(doc, 'Architektur: amd64')
    add_bullet(doc, 'Internetverbindung (für die erstmalige Installation)')

    # --- 2. Installation ---
    doc.add_heading('2. Add-on installieren', level=1)

    doc.add_heading('Schritt 1: Add-on-Repository hinzufügen', level=2)
    add_bullet(doc, 'Home Assistant öffnen')
    add_bullet(doc, 'Einstellungen → Add-ons → Add-on Store (unten rechts)')
    add_bullet(doc, 'Oben rechts: Drei-Punkte-Menü → „Repositories"')
    add_bullet(doc, 'Folgende URL einfügen:')
    add_code(doc, 'https://github.com/DerRegner-DE/ha-device-inventory')
    add_bullet(doc, '„Hinzufügen" klicken, dann Seite neu laden')

    doc.add_heading('Schritt 2: Add-on installieren', level=2)
    add_bullet(doc, 'Im Add-on Store nach „Geräteverwaltung" suchen')
    add_bullet(doc, 'Auf das Add-on klicken')
    add_bullet(doc, '„Installieren" klicken (Download des Docker-Images)')
    add_bullet(doc, 'Warten bis die Installation abgeschlossen ist')

    doc.add_heading('Schritt 3: Add-on konfigurieren', level=2)
    add_bullet(doc, 'Tab „Konfiguration" öffnen')
    add_bullet(doc, 'Sprache wählen: de, en, es, fr, ru')
    add_bullet(doc, 'Optional: Lizenzschlüssel eintragen (für Pro-Features)')
    add_bullet(doc, '„Speichern" klicken')

    doc.add_heading('Schritt 4: Add-on starten', level=2)
    add_bullet(doc, 'Tab „Info" öffnen')
    add_bullet(doc, '„Starten" klicken')
    add_bullet(doc, '„In Seitenleiste anzeigen" aktivieren')
    add_bullet(doc, 'Die App ist jetzt über die HA-Seitenleiste erreichbar')

    add_tip(doc, '„Beim Booten starten" aktivieren, damit das Add-on nach einem HA-Neustart automatisch startet.')

    # --- 3. Konfiguration ---
    doc.add_heading('3. Konfigurationsoptionen', level=1)
    add_table(doc, ['Option', 'Wert', 'Beschreibung'], [
        ['language', 'de / en / es / fr / ru', 'Sprache der Benutzeroberfläche'],
        ['license_key', 'Text (optional)', 'Pro-Lizenzschlüssel für erweiterte Funktionen'],
    ], col_widths=[3, 4, 7])

    # --- 4. Datenspeicherung ---
    doc.add_heading('4. Datenspeicherung', level=1)
    doc.add_paragraph(
        'Alle Daten werden persistent gespeichert und überleben Add-on-Updates:'
    )
    add_table(doc, ['Daten', 'Speicherort', 'Format'], [
        ['Geräte-Datenbank', '/data/db/devices.db', 'SQLite'],
        ['Fotos', '/data/photos/', 'JPEG/PNG/WebP/GIF (max. 10 MB)'],
        ['Lizenz', '/data/db/ (serverseitig)', 'In der Datenbank gespeichert'],
    ], col_widths=[4, 5, 5])

    add_warning(doc, 'Bei einer Deinstallation des Add-ons werden die Daten gelöscht! Vorher ein HA-Backup erstellen.')

    # --- 5. Update ---
    doc.add_heading('5. Add-on aktualisieren', level=1)
    add_bullet(doc, 'Einstellungen → Add-ons → Geräteverwaltung')
    add_bullet(doc, 'Wenn ein Update verfügbar ist, erscheint ein „Aktualisieren"-Button')
    add_bullet(doc, 'Alternativ: Oben rechts Drei-Punkte-Menü → „Nach Updates suchen"')
    add_bullet(doc, '„Aktualisieren" klicken – die Datenbank bleibt erhalten')

    add_tip(doc, 'Add-on-Updates können auch über die HA-Update-Entities automatisiert werden.')

    # --- 6. Fehlerbehebung ---
    doc.add_heading('6. Fehlerbehebung', level=1)

    doc.add_heading('Add-on startet nicht', level=2)
    add_bullet(doc, 'Logs prüfen: Add-on → Tab „Protokoll"')
    add_bullet(doc, 'Port-Konflikte: Sicherstellen, dass Port 3001 nicht belegt ist')
    add_bullet(doc, 'Add-on neu installieren (Daten gehen verloren!)')

    doc.add_heading('App zeigt weiße Seite', level=2)
    add_bullet(doc, 'Browser-Cache leeren (Strg+Umschalt+R)')
    add_bullet(doc, 'Add-on neu starten')
    add_bullet(doc, 'Ingress-URL prüfen (nicht die direkte Port-URL verwenden)')

    doc.add_heading('Kamera funktioniert nicht', level=2)
    add_bullet(doc, '„Take Photo" nutzen (funktioniert immer über Dateiauswahl)')
    add_bullet(doc, '„Use Webcam" benötigt HTTPS oder localhost für getUserMedia')
    add_bullet(doc, 'HA Ingress setzt kein allow="camera" im iframe – getUserMedia kann eingeschränkt sein')
    add_bullet(doc, 'Alternative: App direkt über die HA-URL im Browser öffnen (nicht über Ingress)')

    doc.add_heading('Sync-Symbol rot (keine Verbindung)', level=2)
    add_bullet(doc, 'Backend läuft möglicherweise nicht – Add-on neu starten')
    add_bullet(doc, 'Netzwerkverbindung prüfen')
    add_bullet(doc, 'Im Offline-Modus werden Daten lokal gespeichert und später synchronisiert')

    doc.add_heading('Lizenz wird nicht erkannt', level=2)
    add_bullet(doc, 'Lizenzschlüssel in der Add-on-Konfiguration prüfen')
    add_bullet(doc, 'Add-on nach Eingabe des Schlüssels neu starten')
    add_bullet(doc, 'Schlüssel muss exakt eingegeben werden (Groß-/Kleinschreibung beachten)')

    # --- 7. Deinstallation ---
    doc.add_heading('7. Deinstallation', level=1)
    add_bullet(doc, 'Einstellungen → Add-ons → Geräteverwaltung')
    add_bullet(doc, '„Deinstallieren" klicken')

    add_warning(doc, 'Bei der Deinstallation werden alle Gerätedaten und Fotos gelöscht! Vorher einen Excel-Export oder ein HA-Backup erstellen.')

    # --- 8. Technische Details ---
    doc.add_heading('8. Technische Details', level=1)
    add_table(doc, ['Eigenschaft', 'Wert'], [
        ['GitHub-Repository', 'https://github.com/DerRegner-DE/ha-device-inventory'],
        ['Docker-Image', 'ghcr.io/derregner-de/ha-device-inventory'],
        ['Architektur', 'amd64'],
        ['Interner Port', '3001 (Nginx) → 3002 (FastAPI)'],
        ['Datenbank', 'SQLite (/data/db/devices.db)'],
        ['Frontend-Framework', 'Preact + TypeScript + Tailwind CSS'],
        ['Backend-Framework', 'Python 3.12 + FastAPI + Uvicorn'],
        ['Lizenzmodell', 'Open Core (Free + Pro)'],
    ], col_widths=[5, 9])

    doc.save('Installationsanleitung_Geraeteverwaltung.docx')
    print("Installationsanleitung erstellt.")

# ============================================================
# 3. TESTPLAN
# ============================================================

def create_testplan():
    doc = Document()
    set_style(doc)
    add_title_page(doc, "Geräteverwaltung", "Testplan\nVersion 1.3.5")

    # --- Testumgebung ---
    doc.add_heading('Testumgebung', level=1)
    add_table(doc, ['Eigenschaft', 'Wert'], [
        ['App-Version', '1.3.5'],
        ['Home Assistant', '2026.3.3'],
        ['Testgerät Handy', 'Samsung Galaxy S21 (SM-G996B), Chrome Android'],
        ['Testgerät Desktop', 'Windows 11, Chrome Desktop'],
        ['Zugriff', 'HA-Seitenleiste (Ingress)'],
        ['Add-on', 'Geräteverwaltung v1.3.5'],
    ], col_widths=[5, 9])

    # --- Teil A: Grundfunktionen ---
    doc.add_heading('Teil A: Grundfunktionen (Handy)', level=1)
    doc.add_paragraph('Testgerät: Samsung Galaxy S21, Chrome Android, Zugriff über HA Companion App')

    add_table(doc, ['Nr', 'Testfall', 'Schritte', 'Erwartetes Ergebnis', 'Status'], [
        ['A1', 'App öffnen', 'HA-Seitenleiste → Geräteverwaltung', 'App lädt, Header sichtbar, grüner Punkt', ''],
        ['A2', 'Dashboard leer', 'Dashboard-Tab tippen', 'Zeigt leere Statistik oder Hinweis', ''],
        ['A3', 'Gerät anlegen (Minimal)', 'Hinzufügen → Typ Router → Bezeichnung → Speichern', 'Gerät erscheint in Liste', ''],
        ['A4', 'Gerät anlegen (Voll)', 'Alle Felder ausfüllen inkl. Standort, SN, MAC, IP', 'Alle Felder korrekt gespeichert', ''],
        ['A5', 'Standort wählen', 'Bereich-Dropdown öffnen', '74 Bereiche gruppiert nach Etage', ''],
        ['A6', 'Gerät suchen', 'Geräte → Suchfeld → Text eingeben', 'Nur passende Geräte angezeigt', ''],
        ['A7', 'Filter nach Typ', 'Geräte → Filter-Chip tippen', 'Nur gewählter Typ angezeigt', ''],
        ['A8', 'Gerät bearbeiten', 'Gerät antippen → Bearbeiten', 'Formular vorausgefüllt, speicherbar', ''],
        ['A9', 'Gerät löschen', 'Gerät bearbeiten → Löschen', 'Gerät verschwindet aus Liste', ''],
        ['A10', 'Formular-Validierung', 'Speichern ohne Pflichtfelder', 'Fehlermeldung wird angezeigt', ''],
        ['A11', 'Sonderzeichen', 'Bezeichnung: Küche Öfen Straße', 'Korrekt gespeichert und angezeigt', ''],
        ['A12', 'Dashboard-Statistik', 'Nach mehreren Geräten Dashboard prüfen', 'Korrekte Zählung nach Typ/Standort', ''],
    ], col_widths=[1, 3.5, 4, 4, 1.5])

    # --- Teil B: Kamera & Scanner ---
    doc.add_heading('Teil B: Kamera und Scanner (Pro-Lizenz)', level=1)
    doc.add_paragraph('Voraussetzung: Pro-Lizenz aktiviert')

    add_table(doc, ['Nr', 'Testfall', 'Schritte', 'Erwartetes Ergebnis', 'Status'], [
        ['B1', 'Kamera-Buttons sichtbar', 'Neues Gerät → Foto-Bereich', '„Take Photo" und „Use Webcam" Buttons sichtbar', ''],
        ['B2', 'Take Photo (Handy)', '„Take Photo" antippen', 'Dateiauswahl/Kamera-App öffnet sich', ''],
        ['B3', 'Foto-Vorschau', 'Foto aufnehmen', 'Vorschau mit Bestätigen/Erneut-Buttons', ''],
        ['B4', 'Foto bestätigen', 'Vorschau → Bestätigen', 'Foto wird dem Gerät zugeordnet', ''],
        ['B5', 'Foto erneut aufnehmen', 'Vorschau → Erneut aufnehmen', 'Kamera öffnet sich erneut', ''],
        ['B6', 'Use Webcam (Handy)', '„Use Webcam" antippen', 'Live-Kamera-Stream oder Fallback auf Dateiauswahl', ''],
        ['B7', 'Webcam-Auslöser', 'Live-Stream → Auslöser-Button', 'Foto wird aufgenommen, Vorschau angezeigt', ''],
        ['B8', 'QR-Scanner', 'Scanner-Icon tippen → QR-Code scannen', 'Erkannter Text wird eingetragen', ''],
        ['B9', 'QR Feld-Befüllung', 'QR-Code mit Seriennummer scannen', 'SN-Feld wird befüllt, Abschnitt klappt auf', ''],
        ['B10', 'Barcode scannen', 'Barcode (EAN/MAC) scannen', 'Text korrekt erkannt und eingetragen', ''],
        ['B11', 'Foto hochladen', 'Gerät mit Foto speichern', 'Foto auf Server gespeichert (max. 10 MB)', ''],
        ['B12', 'Foto in Detail', 'Gerät mit Foto öffnen', 'Foto wird in Detailansicht angezeigt', ''],
    ], col_widths=[1, 3.5, 4, 4, 1.5])

    add_warning(doc, 'getUserMedia („Use Webcam") funktioniert möglicherweise nicht über HA Ingress-iframe. Falls B6 fehlschlägt, ist dies eine bekannte Einschränkung. „Take Photo" (B2) sollte immer funktionieren.')

    # --- Teil C: Lizenz & Sprache ---
    doc.add_heading('Teil C: Lizenz und Sprache', level=1)

    add_table(doc, ['Nr', 'Testfall', 'Schritte', 'Erwartetes Ergebnis', 'Status'], [
        ['C1', 'Free-Version Limit', 'Ohne Lizenz 50 Geräte anlegen', 'Ab 51. Gerät Fehlermeldung', ''],
        ['C2', 'Kamera gesperrt (Free)', 'Ohne Lizenz Kamera öffnen', 'Lizenz-Hinweis wird angezeigt', ''],
        ['C3', 'Scanner gesperrt (Free)', 'Ohne Lizenz Scanner öffnen', 'Lizenz-Hinweis wird angezeigt', ''],
        ['C4', 'Export gesperrt (Free)', 'Ohne Lizenz Excel exportieren', 'Lizenz-Hinweis wird angezeigt', ''],
        ['C5', 'Lizenz aktivieren', 'Gültigen Key in Konfiguration eintragen', 'Pro-Features freigeschaltet', ''],
        ['C6', 'Sprache Deutsch', 'language: de in Konfiguration', 'Gesamte UI auf Deutsch', ''],
        ['C7', 'Sprache Englisch', 'language: en in Konfiguration', 'Gesamte UI auf Englisch', ''],
        ['C8', 'Sprache Spanisch', 'language: es in Konfiguration', 'Gesamte UI auf Spanisch', ''],
        ['C9', 'Sprache Französisch', 'language: fr in Konfiguration', 'Gesamte UI auf Französisch', ''],
        ['C10', 'Sprache Russisch', 'language: ru in Konfiguration', 'Gesamte UI auf Russisch', ''],
        ['C11', 'Sprache ohne Lizenz', 'Sprache auf de wechseln ohne Lizenz', 'Bleibt auf Englisch (Lizenz-Einschränkung)', ''],
    ], col_widths=[1, 3.5, 4.5, 3.5, 1.5])

    # --- Teil D: Offline & Sync ---
    doc.add_heading('Teil D: Offline-Modus und Synchronisation', level=1)

    add_table(doc, ['Nr', 'Testfall', 'Schritte', 'Erwartetes Ergebnis', 'Status'], [
        ['D1', 'Offline-Modus', 'WLAN deaktivieren → Gerät anlegen', 'Lokal gespeichert, roter Punkt', ''],
        ['D2', 'Sync nach Offline', 'WLAN wieder aktivieren', 'Gelb → Grün, Daten synchronisiert', ''],
        ['D3', 'Offline-Foto', 'Offline Foto aufnehmen', 'Foto lokal gespeichert', ''],
        ['D4', 'Foto-Sync', 'Online gehen nach Offline-Foto', 'Foto wird hochgeladen', ''],
        ['D5', 'Multi-Device-Sync', 'Handy + PC gleichzeitig öffnen', 'Änderungen werden synchronisiert', ''],
    ], col_widths=[1, 3.5, 4.5, 3.5, 1.5])

    # --- Teil E: Desktop ---
    doc.add_heading('Teil E: Desktop-Browser (Chrome)', level=1)

    add_table(doc, ['Nr', 'Testfall', 'Schritte', 'Erwartetes Ergebnis', 'Status'], [
        ['E1', 'App öffnen', 'HA-Seitenleiste → Geräteverwaltung', 'App lädt, responsive Layout', ''],
        ['E2', 'Gerät anlegen', 'Hinzufügen → Alle Felder → Speichern', 'Gerät wird gespeichert', ''],
        ['E3', 'Kamera (Desktop)', 'Kamera-Icon klicken', 'Webcam angefragt oder Dateiauswahl', ''],
        ['E4', 'Excel-Export', 'Einstellungen → Excel exportieren', 'Excel-Datei wird heruntergeladen', ''],
        ['E5', 'Responsive Design', 'Fenster schmaler ziehen', 'Layout passt sich an', ''],
        ['E6', 'Mehrere Geräte', '5 Geräte hintereinander anlegen', 'Alle gespeichert, Dashboard korrekt', ''],
    ], col_widths=[1, 3.5, 4.5, 3.5, 1.5])

    # --- Teil F: Spezial ---
    doc.add_heading('Teil F: Spezial-Tests', level=1)

    add_table(doc, ['Nr', 'Testfall', 'Schritte', 'Erwartetes Ergebnis', 'Status'], [
        ['F1', 'Add-on-Neustart', 'Add-on stoppen und starten', 'Daten bleiben erhalten', ''],
        ['F2', 'HA-Neustart', 'Home Assistant neu starten', 'Add-on startet automatisch, Daten erhalten', ''],
        ['F3', 'Große Datenmengen', '50+ Geräte anlegen', 'App bleibt performant', ''],
        ['F4', 'MAC-Adresse Format', 'AA:BB:CC:DD:EE:FF und AA-BB-CC-DD-EE-FF eingeben', 'Beide akzeptiert', ''],
        ['F5', 'Lange Texte', 'Langen Text in Anmerkungen eingeben', 'Text gespeichert, scrollbar', ''],
        ['F6', 'Add-on-Update', 'Von 1.3.4 auf 1.3.5 updaten', 'Daten bleiben erhalten, neue Features verfügbar', ''],
    ], col_widths=[1, 3.5, 4.5, 3.5, 1.5])

    # --- Hinweise ---
    doc.add_heading('Hinweise', level=1)
    add_bullet(doc, 'Geräte, die noch herumliegen und nicht erfasst sind, können als Testdaten verwendet werden')
    add_bullet(doc, 'Bei Fehlern: Screenshot machen und Testfall-Nr. notieren')
    add_bullet(doc, 'Status-Spalte: OK / FEHLER / SKIP (nicht testbar)')
    add_bullet(doc, 'Pro-Tests (Teil B, C) erfordern einen gültigen Lizenzschlüssel')
    add_bullet(doc, 'Kamera-Tests am besten bei guter Beleuchtung durchführen')

    doc.save('Testplan_Geraeteverwaltung.docx')
    print("Testplan erstellt.")

# ============================================================
# MAIN
# ============================================================

if __name__ == '__main__':
    os.chdir(os.path.dirname(os.path.abspath(__file__)))
    create_benutzerhandbuch()
    create_installationsanleitung()
    create_testplan()
    print("\nAlle 3 Dokumente mit Umlauten erstellt!")
